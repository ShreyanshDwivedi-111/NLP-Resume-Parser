"""
Resume Parser - Extract information from resumes
Supports: PDF, TXT, DOCX, and Images (JPG, PNG) with OCR
Uses NLP processing and pattern matching
"""

import re
import os
from typing import Dict, List
import PyPDF2
from docx import Document
import docx2txt

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("Warning: pytesseract or PIL not available. Image support disabled.")

from nlp_processor import NLPProcessor


class ResumeParser:
    """
    Parse resumes and extract structured information
    Supports: PDF, TXT, DOCX, JPG, PNG
    """
    
    def __init__(self, nlp_processor: NLPProcessor):
        """Initialize with NLP processor - NO HARDCODED SKILLS!"""
        self.nlp = nlp_processor
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
        except Exception as e:
            print(f"PDF extraction error: {e}")
        return text
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file with multiple encoding fallbacks"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    if content.strip():  # Check if content is not empty
                        print(f"Successfully read TXT file with {encoding} encoding")
                        return content
            except (UnicodeDecodeError, UnicodeError):
                continue  # Try next encoding
            except Exception as e:
                print(f"TXT extraction error with {encoding}: {e}")
                continue
        
        # If all encodings fail, try reading as binary and decode with errors='ignore'
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                if content.strip():
                    print("Read TXT file with UTF-8 (ignoring errors)")
                    return content
        except Exception as e:
            print(f"Final TXT extraction attempt failed: {e}")
        
        return ""  # Return empty string if all attempts fail
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file including tables"""
        try:
            doc = Document(file_path)
            text = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            extracted_text = '\n'.join(text)
            print(f"DOCX: Extracted {len(extracted_text)} characters from {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            
            if not extracted_text:
                print("WARNING: No text extracted from DOCX file!")
                raise ValueError("Empty document - no text content found")
            
            return extracted_text
            
        except Exception as e:
            error_msg = f"DOCX extraction failed: {type(e).__name__}: {str(e)}"
            print(error_msg)
            raise ValueError(error_msg)
    
    def extract_text_from_doc(self, file_path: str) -> str:
        """Extract text from old .doc format (Word 97-2003)"""
        
        # Method 1: Try docx2txt first (works for some .doc files)
        try:
            text = docx2txt.process(file_path)
            if text and text.strip():
                return text
        except Exception:
            pass
        
        # Method 2: Try using Word COM automation (Windows only)
        try:
            import win32com.client
            
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # Open document with absolute path
            doc = word.Documents.Open(os.path.abspath(file_path))
            text = doc.Content.Text
            
            # Close document and Word
            doc.Close(False)
            word.Quit()
            
            if text and text.strip():
                return text
        except ImportError:
            pass
        except Exception:
            try:
                word.Quit()
            except:
                pass
        
        # If both methods fail, give clear instructions
        raise ValueError(
            "Unable to extract text from .doc file. "
            "Please convert to .docx format: "
            "Open in Word → File → Save As → Word Document (*.docx)"
        )
    
    def extract_text_from_image(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        if not OCR_AVAILABLE:
            return "OCR not available. Install pytesseract and PIL."
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            print(f"OCR extraction error: {e}")
            return ""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text based on file type"""
        file_lower = file_path.lower()
        
        if file_lower.endswith('.pdf'):
            return self.extract_text_from_pdf(file_path)
        elif file_lower.endswith('.txt'):
            return self.extract_text_from_txt(file_path)
        elif file_lower.endswith('.docx'):
            return self.extract_text_from_docx(file_path)
        elif file_lower.endswith('.doc'):
            return self.extract_text_from_doc(file_path)
        elif file_lower.endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff')):
            if not OCR_AVAILABLE:
                raise ValueError("Image support requires pytesseract installation")
            return self.extract_text_from_image(file_path)
        else:
            raise ValueError(
                f"Unsupported file format. "
                f"Supported formats: PDF, TXT, DOC, DOCX, JPG, PNG"
            )
    
    def extract_email(self, text: str) -> str:
        """Extract email using regex"""
        pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        match = re.search(pattern, text)
        return match.group(0) if match else ""
    
    def extract_phone(self, text: str) -> str:
        """Extract phone number using regex - multiple formats supported"""
        patterns = [
            r'\(?\d{3}\)?\s*[-.\s]?\d{3}[-.\s]?\d{4}',  # (555) 123-4567 or (555)123-4567 or 555-123-4567
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # +1-555-123-4567
            r'\d{10}',  # 5551234567
            r'\d{3}[-.\s]\d{3}[-.\s]\d{4}'  # 555-123-4567
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        return ""
    
    def extract_skills(self, text: str) -> List[str]:
        """
        Extract skills using improved NLP + technical term preservation
        Captures both processed and original technical terms
        """
        skills = []
        
        # Step 1: Extract keywords using NLP (TF-based importance)
        keywords = self.nlp.extract_keywords(text, top_n=30)
        skills.extend(keywords)
        
        # Step 2: Capture technical terms with special characters (e.g., Node.js, C++, C#)
        # Look for capitalized words and technical patterns
        technical_patterns = [
            r'\b[A-Z][a-zA-Z]*\.js\b',  # Node.js, Next.js, Vue.js
            r'\b[A-Z][a-zA-Z]+\+\+\b',  # C++
            r'\b[A-Z]#\b',              # C#, F#
            r'\b[A-Z][a-zA-Z]+Script\b', # JavaScript, TypeScript
            r'\b[A-Z]{2,}\b',           # AWS, SQL, API, HTML, CSS
            r'\b[A-Z][a-z]+(?:[A-Z][a-z]+)+\b',  # CamelCase: PostgreSQL, MongoDB
        ]
        
        for pattern in technical_patterns:
            matches = re.findall(pattern, text)
            skills.extend([m.lower() for m in matches])
        
        # Step 3: Extract common multi-word technical terms
        multi_word_patterns = [
            r'\b(?:machine learning|deep learning|data science|cloud computing|'
            r'web development|full stack|front end|back end|software engineering|'
            r'ci/cd|devops|microservices|rest api|graphql)\b'
        ]
        
        for pattern in multi_word_patterns:
            matches = re.findall(pattern, text.lower())
            skills.extend(matches)
        
        # Step 4: Tokenize and find frequently mentioned terms
        tokens = self.nlp.tokenize(text.lower())
        lemmatized = self.nlp.lemmatize(text.lower())
        
        for token in set(lemmatized):
            if len(token) > 2 and token not in skills:
                # Add if appears multiple times
                if text.lower().count(token) >= 2:
                    skills.append(token)
        
        # Step 5: Clean and deduplicate
        unique_skills = []
        seen = set()
        for skill in skills:
            skill_clean = skill.strip().lower()
            if skill_clean and len(skill_clean) > 1 and skill_clean not in seen:
                # Filter out common words that aren't skills
                common_words = {'the', 'and', 'for', 'with', 'from', 'this', 'that', 
                               'have', 'has', 'had', 'been', 'were', 'are', 'was'}
                if skill_clean not in common_words:
                    unique_skills.append(skill_clean)
                    seen.add(skill_clean)
        
        return unique_skills[:40]  # Return top 40 skills
    
    def extract_experience(self, text: str) -> str:
        """
        Extract years of experience
        Looks for patterns like "5 years experience", "3+ years"
        """
        patterns = [
            r'(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?\s+experience',
            r'experience\s*:?\s*(\d+)\+?\s*(?:years?|yrs?)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text.lower())
            if match:
                return f"{match.group(1)} years"
        
        return "Not specified"
    
    def parse(self, file_path: str) -> Dict:
        """
        Parse resume and extract all information
        
        Returns:
            Dictionary with extracted data:
            - email, phone, skills, experience, keywords
        """
        try:
            # Extract text
            text = self.extract_text(file_path)
            
            if not text:
                return {"error": "Could not extract text from file"}
            
            # Extract information
            email = self.extract_email(text)
            phone = self.extract_phone(text)
            skills = self.extract_skills(text)
            experience = self.extract_experience(text)
            
            # Extract keywords using NLP
            keywords = self.nlp.extract_keywords(text, top_n=15)
            
            return {
                "email": email,
                "phone": phone,
                "skills": skills,
                "experience": experience,
                "keywords": keywords,
                "raw_text": text[:500]  # First 500 chars for reference
            }
        
        except Exception as e:
            return {"error": str(e)}
